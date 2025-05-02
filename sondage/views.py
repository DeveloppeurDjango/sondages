from django.shortcuts import render, redirect
from django.contrib import messages
from django.contrib.auth.decorators import user_passes_test, login_required
from django.contrib.auth import login, logout, authenticate
from django.contrib.auth.forms import AuthenticationForm
from django.contrib.auth.models import User
from .models import Section, Question, Option, Response, ResponseOption
from django.db.models import Count, Avg
import os
from django.conf import settings
from django.http import HttpResponse
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

def is_admin(user):
    return user.is_staff

def login_view(request):
    if request.method == 'POST':
        form = AuthenticationForm(request, data=request.POST)
        if form.is_valid():
            username = form.cleaned_data.get('username')
            password = form.cleaned_data.get('password')
            user = authenticate(username=username, password=password)
            if user is not None:
                login(request, user)
                messages.success(request, f"Bienvenue {username}!")
                return redirect('sondage:survey_form')
            else:
                messages.error(request, "Nom d'utilisateur ou mot de passe invalide.")
        else:
            messages.error(request, "Nom d'utilisateur ou mot de passe invalide.")
    else:
        form = AuthenticationForm()
    return render(request, 'sondage/login.html', {'form': form})

def register_view(request):
    if request.method == 'POST':
        username = request.POST.get('username')
        first_name = request.POST.get('first_name')
        last_name = request.POST.get('last_name')
        password1 = request.POST.get('password1')
        password2 = request.POST.get('password2')

        if password1 != password2:
            messages.error(request, "Les mots de passe ne correspondent pas.")
            return render(request, 'sondage/register.html')

        if User.objects.filter(username=username).exists():
            messages.error(request, "Ce nom d'utilisateur est déjà pris.")
            return render(request, 'sondage/register.html')

        user = User.objects.create_user(
            username=username,
            first_name=first_name,
            last_name=last_name,
            password=password1
        )
        messages.success(request, "Inscription réussie ! Veuillez vous connecter.")
        return redirect('sondage:login')

    return render(request, 'sondage/register.html')

def logout_view(request):
    logout(request)
    messages.info(request, "Vous avez été déconnecté.")
    return redirect('sondage:login')

@login_required
def survey_form(request):
    sections = Section.objects.all().prefetch_related('questions__options')
    return render(request, 'sondage/survey_form.html', {'sections': sections})

@login_required
def submit_survey(request):
    if request.method == 'POST':
        response = Response.objects.create(user=request.user)
        
        for question in Question.objects.all():
            if question.is_text:
                field_name = f'question_{question.id}'
                if field_name in request.POST:
                    text_response = request.POST.get(field_name)
                    ResponseOption.objects.create(
                        response=response,
                        question=question,
                        text_response=text_response
                    )
            elif question.is_likert:
                for option in question.options.all():
                    field_name = f'question_{question.id}_{option.id}'
                    if field_name in request.POST:
                        likert_value = int(request.POST.get(field_name))
                        ResponseOption.objects.create(
                            response=response,
                            question=question,
                            option=option,
                            likert_value=likert_value
                        )
            else:
                field_name = f'question_{question.id}'
                if field_name in request.POST:
                    option_id = request.POST.get(field_name)
                    try:
                        option = Option.objects.get(id=option_id, question=question)
                        ResponseOption.objects.create(
                            response=response,
                            question=question,
                            option=option
                        )
                    except Option.DoesNotExist:
                        messages.error(request, f"Option invalide pour la question {question.number}")
                        return redirect('sondage:survey_form')
        
        messages.success(request, "Merci pour votre participation !")
        return render(request, 'sondage/thank_you.html')
    
    return redirect('sondage:survey_form')

@login_required
def generate_report(request):
    sections = Section.objects.all().prefetch_related('questions__options')
    report_data = []
    
    for section in sections:
        section_data = {
            'name': section.name,
            'questions': []
        }
        
        for question in section.questions.all():
            if question.is_text:
                responses = ResponseOption.objects.filter(question=question).values_list('text_response', flat=True)
                question_data = {
                    'question': question,
                    'type': 'text',
                    'responses': list(responses)
                }
            elif question.is_likert:
                options_data = []
                for option in question.options.all():
                    responses = ResponseOption.objects.filter(
                        question=question, 
                        option=option
                    ).values('likert_value').annotate(count=Count('likert_value')).order_by('likert_value')
                    
                    scale_counts = {str(i): 0 for i in range(0, 6)}  # De 0 à 5
                    for response in responses:
                        if response['likert_value'] is not None:
                            scale_counts[str(response['likert_value'])] = response['count']
                    
                    total = sum(int(value) * count for value, count in scale_counts.items())
                    count = sum(scale_counts.values())
                    average = round(total / count, 2) if count > 0 else 0
                    
                    options_data.append({
                        'option': option.text,
                        'average': average,
                        'scale_counts': scale_counts,
                        'total_responses': count
                    })
                
                question_data = {
                    'question': question,
                    'type': 'likert',
                    'options_data': options_data
                }
            else:
                responses = ResponseOption.objects.filter(question=question).values('option__text').annotate(
                    count=Count('option')
                ).order_by('-count')
                
                total = sum(r['count'] for r in responses)
                question_data = {
                    'question': question,
                    'type': 'multiple_choice',
                    'responses': [
                        {
                            'option': r['option__text'],
                            'count': r['count'],
                            'percentage': round((r['count'] / total * 100) if total > 0 else 0, 1)
                        }
                        for r in responses
                    ]
                }
            
            section_data['questions'].append(question_data)
        
        report_data.append(section_data)

    return render(request, 'sondage/report.html', {
        'sections': report_data
    })

@login_required
@user_passes_test(is_admin)
def admin_stats(request):
    responses = Response.objects.select_related('user').prefetch_related(
        'selected_options',
        'selected_options__question'
    ).order_by('-created_at')
    
    sections = Section.objects.all().prefetch_related('questions__options')
    report_data = []
    
    for section in sections:
        section_data = {
            'name': section.name,
            'questions': []
        }
        
        for question in section.questions.all():
            if question.is_text:
                responses = ResponseOption.objects.filter(question=question).values_list('text_response', flat=True)
                question_data = {
                    'question': question,
                    'type': 'text',
                    'responses': list(responses)
                }
            elif question.is_likert:
                options_data = []
                for option in question.options.all():
                    responses = ResponseOption.objects.filter(
                        question=question, 
                        option=option
                    ).values('likert_value').annotate(count=Count('likert_value')).order_by('likert_value')
                    
                    scale_counts = {str(i): 0 for i in range(0, 6)}  # De 0 à 5
                    for response in responses:
                        if response['likert_value'] is not None:
                            scale_counts[str(response['likert_value'])] = response['count']
                    
                    total = sum(int(value) * count for value, count in scale_counts.items())
                    count = sum(scale_counts.values())
                    average = round(total / count, 2) if count > 0 else 0
                    
                    options_data.append({
                        'option': option.text,
                        'average': average,
                        'scale_counts': scale_counts,
                        'total_responses': count
                    })
                
                question_data = {
                    'question': question,
                    'type': 'likert',
                    'options_data': options_data
                }
            else:
                responses = ResponseOption.objects.filter(question=question).values('option__text').annotate(
                    count=Count('option')
                ).order_by('-count')
                
                total = sum(r['count'] for r in responses)
                question_data = {
                    'question': question,
                    'type': 'multiple_choice',
                    'responses': [
                        {
                            'option': r['option__text'],
                            'count': r['count'],
                            'percentage': round((r['count'] / total * 100) if total > 0 else 0, 1)
                        }
                        for r in responses
                    ]
                }
            
            section_data['questions'].append(question_data)
        
        report_data.append(section_data)

    return render(request, 'sondage/admin_stats.html', {
        'responses': responses,
        'sections': report_data
    })

@login_required
@user_passes_test(is_admin)
def admin_stats_word(request):
    sections = Section.objects.all().prefetch_related('questions__options')
    doc = Document()
    
    # Titre du document
    title = doc.add_heading('Rapport des statistiques', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for section in sections:
        # Titre de la section
        doc.add_heading(section.name, level=1)
        
        for question in section.questions.all():
            # Titre de la question
            doc.add_heading(question.text, level=2)
            
            if question.is_text:
                responses = ResponseOption.objects.filter(question=question).values_list('text_response', flat=True)
                for response in responses:
                    doc.add_paragraph(response)
            
            elif question.is_likert:
                table = doc.add_table(rows=1, cols=3)
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Option'
                hdr_cells[1].text = 'Moyenne'
                hdr_cells[2].text = 'Total réponses'
                
                for option in question.options.all():
                    responses = ResponseOption.objects.filter(
                        question=question, 
                        option=option
                    ).values('likert_value').annotate(count=Count('likert_value'))
                    
                    scale_counts = {str(i): 0 for i in range(0, 6)}  # De 0 à 5
                    for response in responses:
                        if response['likert_value'] is not None:
                            scale_counts[str(response['likert_value'])] = response['count']
                    
                    total = sum(r['count'] for r in responses)
                    sum_values = sum(r['likert_value'] * r['count'] for r in responses)
                    average = round(sum_values / total, 2) if total > 0 else 0
                    
                    row_cells = table.add_row().cells
                    row_cells[0].text = option.text
                    row_cells[1].text = str(average)
                    row_cells[2].text = str(total)
            
            else:
                table = doc.add_table(rows=1, cols=3)
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Option'
                hdr_cells[1].text = 'Nombre'
                hdr_cells[2].text = 'Pourcentage'
                
                responses = ResponseOption.objects.filter(question=question).values('option__text').annotate(
                    count=Count('option')
                ).order_by('-count')
                
                total = sum(r['count'] for r in responses)
                for r in responses:
                    row_cells = table.add_row().cells
                    row_cells[0].text = r['option__text']
                    row_cells[1].text = str(r['count'])
                    row_cells[2].text = f"{round((r['count'] / total * 100) if total > 0 else 0, 1)}%"
    
    # Sauvegarder le document
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=statistiques.docx'
    doc.save(response)
    return response

@login_required
@user_passes_test(is_admin)
def admin_stats_pdf(request):
    sections = Section.objects.all().prefetch_related('questions__options')
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    elements = []
    styles = getSampleStyleSheet()
    
    # Titre du document
    title = Paragraph("Rapport des statistiques", styles['Title'])
    elements.append(title)
    elements.append(Spacer(1, 12))
    
    for section in sections:
        # Titre de la section
        section_title = Paragraph(section.name, styles['Heading1'])
        elements.append(section_title)
        elements.append(Spacer(1, 12))
        
        for question in section.questions.all():
            # Titre de la question
            question_title = Paragraph(question.text, styles['Heading2'])
            elements.append(question_title)
            elements.append(Spacer(1, 12))
            
            if question.is_text:
                responses = ResponseOption.objects.filter(question=question).values_list('text_response', flat=True)
                for response in responses:
                    elements.append(Paragraph(response, styles['Normal']))
                    elements.append(Spacer(1, 6))
            
            elif question.is_likert:
                data = [['Option', 'Moyenne', 'Total réponses']]
                for option in question.options.all():
                    responses = ResponseOption.objects.filter(
                        question=question, 
                        option=option
                    ).values('likert_value').annotate(count=Count('likert_value'))
                    
                    scale_counts = {str(i): 0 for i in range(0, 6)}  # De 0 à 5
                    for response in responses:
                        if response['likert_value'] is not None:
                            scale_counts[str(response['likert_value'])] = response['count']
                    
                    total = sum(r['count'] for r in responses)
                    sum_values = sum(r['likert_value'] * r['count'] for r in responses)
                    average = round(sum_values / total, 2) if total > 0 else 0
                    
                    data.append([option.text, str(average), str(total)])
                
                table = Table(data)
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 14),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
                    ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                    ('FONTSIZE', (0, 1), (-1, -1), 12),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                elements.append(table)
                elements.append(Spacer(1, 12))
            
            else:
                data = [['Option', 'Nombre', 'Pourcentage']]
                responses = ResponseOption.objects.filter(question=question).values('option__text').annotate(
                    count=Count('option')
                ).order_by('-count')
                
                total = sum(r['count'] for r in responses)
                for r in responses:
                    percentage = round((r['count'] / total * 100) if total > 0 else 0, 1)
                    data.append([r['option__text'], str(r['count']), f"{percentage}%"])
                
                table = Table(data)
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 14),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
                    ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                    ('FONTSIZE', (0, 1), (-1, -1), 12),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                elements.append(table)
                elements.append(Spacer(1, 12))
    
    doc.build(elements)
    pdf = buffer.getvalue()
    buffer.close()
    
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename=statistiques.pdf'
    response.write(pdf)
    return response
